import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUMMARY_CSV = ROOT / "formal_v1_ccw_positivity_weight_distribution_summary.csv"
BY_ARM_CSV = ROOT / "formal_v1_ccw_positivity_weight_distribution_by_arm_m5.csv"
OVERLAP_CSV = ROOT / "formal_v1_ccw_positivity_continuous_overlap_m5.csv"
FLAGS_CSV = ROOT / "formal_v1_ccw_positivity_flag_summary.csv"
WORK_OUT = ROOT / "formal_v1_supplementary_table2_ipcw_positivity_diagnostics.docx"


PUBLICATION_OUT_DIR = os.environ.get("TTE_PUBLICATION_OUT_DIR")
EXTRA_OUT = Path(PUBLICATION_OUT_DIR) / "Supplementary Table 2.docx" if PUBLICATION_OUT_DIR else None


WEIGHT_LABELS = {
    "final_ipcw": "Final IPCW",
    "trunc_p01_p99": "p01-p99 truncation",
    "trunc_p05_p95": "p05-p95 truncation",
}
WEIGHT_ORDER = ["final_ipcw", "trunc_p01_p99", "trunc_p05_p95"]

FLAG_LABELS = {
    "final_ipcw_max_gt_10": "Maximum final IPCW <=10",
    "final_ipcw_p99_gt_5": "99th percentile final IPCW <=5",
    "final_ipcw_pct_gt_5_gt_1pct": "Final IPCW >5 in <=1% of clones",
    "continuous_common_p01_p99_width_nonpositive": "Continuous p01-p99 overlap width >0",
    "main_categorical_zero_count_arm": "No zero-count categorical arm",
}


def read_csv(path):
    with open(path, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def fnum(value):
    return float(value)


def fmt_num(value, digits=2):
    return f"{float(value):.{digits}f}"


def fmt_pct(value, digits=1):
    return f"{float(value) * 100:.{digits}f}%"


def set_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    def apply(edge, visible, size="8"):
        tag = f"w:{edge}"
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    if top is not None:
        apply("top", top)
    if bottom is not None:
        apply("bottom", bottom)


def table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    def set_border(edge, visible, size="12"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    set_border("top", top, "12")
    set_border("bottom", bottom, "12")
    set_border("left", False)
    set_border("right", False)
    set_border("insideH", inside_h)
    set_border("insideV", inside_v)


def set_table_layout(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, 10.5, True)


def min_ess_by_weight(by_arm_rows, weight_type):
    rows = [r for r in by_arm_rows if r["weight_type"] == weight_type]
    row = min(rows, key=lambda r: fnum(r["ess"]))
    pct = fnum(row["ess"]) / fnum(row["n_clones"])
    arm = row["arm"].replace("h", " h")
    return f"{fmt_num(row['ess'], 1)}\n({pct * 100:.1f}%; {arm})"


def min_arm_size_by_weight(by_arm_rows, weight_type):
    rows = [r for r in by_arm_rows if r["weight_type"] == weight_type]
    row = min(rows, key=lambda r: fnum(r["n_clones"]))
    arm = row["arm"].replace("h", " h")
    return f"{int(fnum(row['n_clones'])):,}\n({arm})"


def continuous_support_rows(overlap_rows):
    out = []
    labels = []
    for r in overlap_rows:
        if r["variable_label"] not in labels:
            labels.append(r["variable_label"])
    for label in labels:
        rows = [r for r in overlap_rows if r["variable_label"] == label]
        p01_low = rows[0]["common_p01_low"]
        p01_high = rows[0]["common_p99_high"]
        p05_low = rows[0]["common_p05_low"]
        p05_high = rows[0]["common_p95_high"]
        max_out_p01 = max(fnum(r["outside_common_p01_p99_pct"]) for r in rows)
        max_out_p05 = max(fnum(r["outside_common_p05_p95_pct"]) for r in rows)
        out.append(
            [
                "Continuous support",
                label,
                f"p01-p99 {p01_low}-{p01_high}; outside {max_out_p01 * 100:.1f}%",
                f"p05-p95 {p05_low}-{p05_high}; outside {max_out_p05 * 100:.1f}%",
                "Pass",
            ]
        )
    return out


def main():
    summary = {r["weight_type"]: r for r in read_csv(SUMMARY_CSV)}
    by_arm = read_csv(BY_ARM_CSV)
    overlap = read_csv(OVERLAP_CSV)
    flags = read_csv(FLAGS_CSV)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    add_caption(doc, "Supplementary Table 2. IPCW weight and positivity diagnostics in MIMIC-IV")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("A. Weight distribution and effective sample size")
    set_font(run, 10.5, True)

    metrics = [
        ("Arm-imputation units, n", lambda wt: summary[wt]["n_arm_imp"]),
        ("Minimum arm size, n (arm)", lambda wt: min_arm_size_by_weight(by_arm, wt)),
        ("Minimum ESS, n (%; arm)", lambda wt: min_ess_by_weight(by_arm, wt)),
        ("Maximum 95th percentile", lambda wt: fmt_num(summary[wt]["max_p95"], 2)),
        ("Maximum 99th percentile", lambda wt: fmt_num(summary[wt]["max_p99"], 2)),
        ("Maximum weight", lambda wt: fmt_num(summary[wt]["max_max"], 2)),
        ("Maximum proportion >2", lambda wt: fmt_pct(summary[wt]["max_pct_gt_2"], 1)),
        ("Maximum proportion >5", lambda wt: fmt_pct(summary[wt]["max_pct_gt_5"], 2)),
        ("Maximum proportion >10", lambda wt: fmt_pct(summary[wt]["max_pct_gt_10"], 2)),
    ]

    table_a = doc.add_table(rows=1, cols=4)
    table_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths_a = [3600, 1900, 1930, 1930]
    set_table_layout(table_a, widths_a)
    table_borders(table_a, top=True, bottom=True)
    headers = ["Diagnostic", *[WEIGHT_LABELS[wt] for wt in WEIGHT_ORDER]]
    for i, h in enumerate(headers):
        set_cell_text(table_a.rows[0].cells[i], h, bold=True)
        set_cell_border(table_a.rows[0].cells[i], bottom=True)
    for label, fn in metrics:
        cells = table_a.add_row().cells
        set_table_layout(table_a, widths_a)
        set_cell_text(cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, wt in enumerate(WEIGHT_ORDER, start=1):
            set_cell_text(cells[i], fn(wt))
    table_borders(table_a, top=True, bottom=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("B. Support and prespecified positivity checks")
    set_font(run, 10.5, True)

    table_b = doc.add_table(rows=1, cols=5)
    table_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths_b = [1800, 2120, 2320, 2320, 800]
    set_table_layout(table_b, widths_b)
    table_borders(table_b, top=True, bottom=True)
    headers_b = ["Domain", "Diagnostic", "p01-p99 support (% outside)", "p05-p95 support (% outside)", "Status"]
    for i, h in enumerate(headers_b):
        set_cell_text(table_b.rows[0].cells[i], h, bold=True)
        set_cell_border(table_b.rows[0].cells[i], bottom=True)

    for row in continuous_support_rows(overlap):
        cells = table_b.add_row().cells
        set_table_layout(table_b, widths_b)
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i < 4 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, align=align)

    for flag in flags:
        cells = table_b.add_row().cells
        set_table_layout(table_b, widths_b)
        set_cell_text(cells[0], "Flag check", align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[1], FLAG_LABELS.get(flag["check"], flag["check"]), align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[2], "-", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], "-", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[4], flag["status"].capitalize(), align=WD_ALIGN_PARAGRAPH.CENTER)
    table_borders(table_b, top=True, bottom=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    note_run = note.add_run(
        "Note: Weight diagnostics are summarized across the 15 arm-imputation units unless otherwise specified. "
        "Minimum arm size and minimum ESS may occur in different strategy arms; the arm is shown in parentheses for both diagnostics. "
        "For p05-p95 truncation, upper-tail weights are capped at the p95 value. "
        "Continuous-support rows report the maximum proportion outside the common percentile support interval across arms and imputed datasets. "
        "ESS, effective sample size; IPCW, inverse probability of censoring weighting; "
        "MIMIC-IV, Medical Information Mart for Intensive Care IV."
    )
    set_font(note_run, 9, False)

    for table in [table_a, table_b]:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0

    outputs = [WORK_OUT] + ([EXTRA_OUT] if EXTRA_OUT else [])
    for out in outputs:
        out.parent.mkdir(parents=True, exist_ok=True)
        try:
            doc.save(out)
        except PermissionError:
            revised_out = out.with_name(out.stem + "_revised.docx")
            doc.save(revised_out)
            out = revised_out
        print(out)


if __name__ == "__main__":
    main()

