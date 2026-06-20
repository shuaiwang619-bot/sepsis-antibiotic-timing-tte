from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


base = Path(__file__).resolve().parents[1]
csv_path = base / "formal_v1_table1_baseline_characteristics_m5.csv"
out_path = base / "formal_v1_table1_baseline_characteristics_three_line.docx"

with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
    rows = list(csv.DictReader(f))

cols = [
    "characteristic",
    "unweighted_0_3h",
    "unweighted_3_6h",
    "unweighted_6_12h",
    "max_smd_before",
    "weighted_0_3h",
    "weighted_3_6h",
    "weighted_6_12h",
    "max_smd_after",
]
headers1 = [
    "Characteristic",
    "Unweighted",
    "",
    "",
    "Before",
    "IPCW weighted",
    "",
    "",
    "After",
]
headers2 = ["", "0-3h", "3-6h", "6-12h", "Max SMD", "0-3h", "3-6h", "6-12h", "Max SMD"]
widths = [1.75, 1.05, 1.05, 1.05, 0.62, 1.15, 1.15, 1.15, 0.62]


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=35, start=60, bottom=35, end=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, val in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        if val is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(val.get("sz", 8)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), val.get("color", "000000"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def fill_cell(cell, text, bold=False, italic=False, align="center", size=7.2, color="000000"):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.35)
section.bottom_margin = Inches(0.35)
section.left_margin = Inches(0.45)
section.right_margin = Inches(0.45)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(9)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
title_run = title.add_run("Table 1. Baseline characteristics of strategy-adherent clones before and after IPCW")
title_run.bold = True
title_run.font.size = Pt(9.8)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(3)
subtitle_run = subtitle.add_run(
    "Continuous variables are shown as median [IQR]; categorical variables are shown as n (%). "
    "Weighted columns show IPCW-weighted pseudo-counts and percentages. SMD is the maximum "
    "absolute pairwise standardized mean difference across the three antibiotic initiation strategies."
)
subtitle_run.font.size = Pt(7.2)
subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

table = doc.add_table(rows=len(rows) + 2, cols=len(cols))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

for row in table.rows:
    for j, cell in enumerate(row.cells):
        set_cell_width(cell, widths[j])
        set_cell_margins(cell)
        set_cell_border(cell)

for j, text in enumerate(headers1):
    fill_cell(table.cell(0, j), text, bold=True, size=6.9)
for j, text in enumerate(headers2):
    fill_cell(table.cell(1, j), text, bold=True, size=6.9)

table.cell(0, 1).merge(table.cell(0, 3))
table.cell(0, 5).merge(table.cell(0, 7))
set_repeat_table_header(table.rows[0])
set_repeat_table_header(table.rows[1])

for cell in table.rows[0].cells:
    set_cell_border(cell, top={"sz": 12}, bottom=None)
for cell in table.rows[1].cells:
    set_cell_border(cell, top=None, bottom={"sz": 8})

for i, row_data in enumerate(rows, start=2):
    row = table.rows[i]
    if row_data["row_type"] == "section":
        merged = row.cells[0]
        for k in range(1, len(cols)):
            merged = merged.merge(row.cells[k])
        fill_cell(row.cells[0], row_data["characteristic"], bold=True, align="left", size=6.9)
        set_cell_border(row.cells[0])
        continue

    for j, col in enumerate(cols):
        align = "left" if j == 0 else "center"
        fill_cell(row.cells[j], row_data[col], align=align, size=6.7)
        set_cell_border(row.cells[j])

for cell in table.rows[-1].cells:
    set_cell_border(cell, bottom={"sz": 12})

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(3)
note.paragraph_format.space_after = Pt(0)
note_run = note.add_run(
    "Notes: The table uses m=5 imputed baseline datasets; per-imputation summaries were averaged "
    "for display. Weighted columns use final stabilized IPCW among strategy-adherent clones. The SOFA "
    "missing indicator is sparse but retained because it is included in the model. Mean BP variables "
    "are auxiliary/sensitivity covariates rather than primary model covariates."
)
note_run.font.size = Pt(6.8)
note_run.font.color.rgb = RGBColor(80, 80, 80)

doc.save(out_path)
print(out_path)
