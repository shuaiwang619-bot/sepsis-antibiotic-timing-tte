import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "formal_v1_results_sensitivity_table.csv"
WORK_OUT = ROOT / "formal_v1_supplementary_table1_sensitivity_analyses.docx"


PUBLICATION_OUT_DIR = os.environ.get("TTE_PUBLICATION_OUT_DIR")
EXTRA_OUT = Path(PUBLICATION_OUT_DIR) / "Supplementary Table 1.docx" if PUBLICATION_OUT_DIR else None


MODEL_LABELS = {
    "Main: coarse-time pooled logistic": "Main CCW/IPCW pooled logistic",
    "Sensitivity: full hourly time-axis pooled logistic": "Full hourly pooled logistic",
    "Sensitivity: weighted Cox model": "Weighted Cox model",
    "Sensitivity: pooled logistic with p01-p99 weight truncation": "p01-p99 weight truncation",
    "Sensitivity: pooled logistic with p05-p95 weight truncation": "p05-p95 weight truncation",
}

MODEL_ORDER = list(MODEL_LABELS)
CONTRAST_ORDER = ["3-6h vs 0-3h", "6-12h vs 0-3h", "6-12h vs 3-6h"]
CONTRAST_LABELS = {
    "3-6h vs 0-3h": "3-6 h vs 0-3 h",
    "6-12h vs 0-3h": "6-12 h vs 0-3 h",
    "6-12h vs 3-6h": "6-12 h vs 3-6 h",
}


def set_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    def apply(edge, visible, size="8"):
        tag = f"w:{edge}"
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    if top is not None:
        apply("top", top)
    if bottom is not None:
        apply("bottom", bottom)


def set_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = tcPr.first_child_found_in("w:noWrap")
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tcPr.append(no_wrap)


def table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    def set_border(edge, visible, size="12"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    set_border("top", top, "12")
    set_border("bottom", bottom, "12")
    set_border("left", False)
    set_border("right", False)
    set_border("insideH", inside_h)
    set_border("insideV", inside_v)


def set_table_layout(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def normalize_effect(effect):
    return effect.replace("OR ", "OR\n").replace("HR ", "HR\n")


def main():
    with open(CSV_PATH, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))

    lookup = {(r["model"], r["contrast"]): r for r in rows}

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Supplementary Table 1. Primary and sensitivity outcome models in MIMIC-IV")
    set_font(run, 10.5, True)

    table = doc.add_table(rows=2, cols=7)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False)
    widths = [2400, 1550, 770, 1550, 770, 1550, 770]
    set_table_layout(table, widths)

    header0 = table.rows[0].cells
    header1 = table.rows[1].cells
    set_cell_text(header0[0], "Model", bold=True)
    for idx, contrast in enumerate(CONTRAST_ORDER):
        start = 1 + idx * 2
        merged = header0[start].merge(header0[start + 1])
        set_cell_text(merged, CONTRAST_LABELS[contrast], bold=True)
        set_cell_text(header1[start], "Estimate\n(95% CI)", bold=True, size=10)
        set_cell_text(header1[start + 1], "P value", bold=True, size=10)
        set_no_wrap(header1[start + 1])
    set_cell_text(header1[0], "", bold=True)

    for cell in table.rows[1].cells:
        set_cell_border(cell, bottom=True)

    for model in MODEL_ORDER:
        cells = table.add_row().cells
        set_table_layout(table, widths)
        set_cell_text(cells[0], MODEL_LABELS[model], align=WD_ALIGN_PARAGRAPH.LEFT)
        for idx, contrast in enumerate(CONTRAST_ORDER):
            r = lookup[(model, contrast)]
            start = 1 + idx * 2
            set_cell_text(cells[start], normalize_effect(r["effect"]), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
            set_cell_text(cells[start + 1], r["p_value"], align=WD_ALIGN_PARAGRAPH.CENTER)
            set_no_wrap(cells[start + 1])

    table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False)
    for cell in table.rows[1].cells:
        set_cell_border(cell, bottom=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    note_run = note.add_run(
        "Note: The main model was an IPCW-weighted, covariate-adjusted pooled logistic regression model. "
        "Effect estimates are odds ratios except for the weighted Cox model, which reports hazard ratios. "
        "Weight-truncation analyses used final IPCW truncated at the indicated percentiles. "
        "CCW, clone-censor-weight; CI, confidence interval; HR, hazard ratio; IPCW, inverse probability of "
        "censoring weighting; MIMIC-IV, Medical Information Mart for Intensive Care IV; OR, odds ratio."
    )
    set_font(note_run, 9, False)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

    outputs = [WORK_OUT] + ([EXTRA_OUT] if EXTRA_OUT else [])
    for out in outputs:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)

    for out in outputs:
        print(out)


if __name__ == "__main__":
    main()

