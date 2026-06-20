from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


base = Path(__file__).resolve().parents[1]
main_csv = base / "formal_v1_results_main_model_table.csv"
sens_csv = base / "formal_v1_results_sensitivity_table.csv"
forest_png = base / "formal_v1_results_forest_plot.png"
out_path = base / "formal_v1_results_tables_and_forest.docx"


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


main_rows = read_csv(main_csv)
sens_rows = read_csv(sens_csv)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=55, start=80, bottom=55, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, val in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        if val is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(val.get("sz", 8)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), val.get("color", "000000"))


def fill_cell(cell, text, bold=False, align="center", size=8.0, color="000000"):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p


def add_three_line_table(doc, headers, rows, keys, widths, aligns=None, font_size=8.0):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if aligns is None:
        aligns = ["center"] * len(headers)
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            set_cell_width(cell, widths[j])
            set_cell_margins(cell)
            set_cell_border(cell)
    for j, header in enumerate(headers):
        fill_cell(table.cell(0, j), header, bold=True, align=aligns[j], size=font_size)
        set_cell_border(table.cell(0, j), top={"sz": 12}, bottom={"sz": 8})
    for i, row_data in enumerate(rows, start=1):
        for j, key in enumerate(keys):
            fill_cell(table.cell(i, j), row_data.get(key, ""), align=aligns[j], size=font_size)
            set_cell_border(table.cell(i, j))
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 12})
    return table


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(9)

add_title(doc, "Table 2. Main analysis: antibiotic initiation window and 30-day mortality")
add_note(
    doc,
    "Main model is the locked coarse-time weighted pooled logistic model. Risks are standardized 30-day risks; "
    "risk difference is shown in percentage points.",
)
add_three_line_table(
    doc,
    headers=[
        "Contrast",
        "Reference risk",
        "Comparison risk",
        "Risk difference",
        "Risk ratio",
        "OR (95% CI)",
        "P value",
    ],
    rows=main_rows,
    keys=[
        "contrast",
        "reference_30d_risk",
        "comparison_30d_risk",
        "risk_difference",
        "risk_ratio",
        "odds_ratio_95_ci",
        "p_value",
    ],
    widths=[1.7, 1.25, 1.35, 1.35, 0.9, 1.35, 0.8],
    aligns=["left", "center", "center", "center", "center", "center", "center"],
    font_size=8.2,
)

add_title(doc, "Table 3. Sensitivity analyses")
add_note(
    doc,
    "Full hourly time-axis, weighted Cox, and weight-truncation analyses are shown as robustness checks. "
    "OR denotes odds ratio; HR denotes hazard ratio.",
)
short_model = {
    "Main: coarse-time pooled logistic": "Main coarse-time",
    "Sensitivity: full hourly time-axis pooled logistic": "Full hourly pooled logistic",
    "Sensitivity: weighted Cox model": "Weighted Cox",
    "Sensitivity: pooled logistic with p01-p99 weight truncation": "Truncation p01-p99",
    "Sensitivity: pooled logistic with p05-p95 weight truncation": "Truncation p05-p95",
}
sens_rows_short = []
for row in sens_rows:
    nr = dict(row)
    nr["model"] = short_model.get(row["model"], row["model"])
    sens_rows_short.append(nr)

add_three_line_table(
    doc,
    headers=["Model", "Contrast", "Effect estimate (95% CI)", "P value"],
    rows=sens_rows_short,
    keys=["model", "contrast", "effect", "p_value"],
    widths=[2.25, 1.55, 1.75, 0.8],
    aligns=["left", "left", "center", "center"],
    font_size=7.8,
)

doc.add_page_break()
add_title(doc, "Figure 1. Sensitivity forest plot", size=12)
if forest_png.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(forest_png), width=Inches(6.8))

doc.save(out_path)
print(out_path)
