import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "formal_v1_results_main_model_table.csv"
WORK_OUT = ROOT / "formal_v1_table2_primary_outcome_estimates.docx"


PUBLICATION_OUT_DIR = os.environ.get("TTE_PUBLICATION_OUT_DIR")
EXTRA_OUT = Path(PUBLICATION_OUT_DIR) / "Table 2.docx" if PUBLICATION_OUT_DIR else None


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_font(paragraph, size=10.5, bold=False):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    def set_border(edge, visible, size="12"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    set_border("top", top, "12")
    set_border("bottom", bottom, "12")
    set_border("left", False)
    set_border("right", False)
    set_border("insideH", inside_h)
    set_border("insideV", inside_v)


def set_cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    def apply(edge, visible, size="8"):
        tag = f"w:{edge}"
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        if visible:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")

    if top is not None:
        apply("top", top)
    if bottom is not None:
        apply("bottom", bottom)


def set_table_layout(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def main():
    with open(CSV_PATH, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Table 2. Primary MIMIC-IV outcome estimates")
    title_run.bold = True
    set_paragraph_font(title, 10.5, True)
    keep_with_next(title)

    headers = [
        "Contrast",
        "Reference risk (%)",
        "Comparison risk (%)",
        "Risk difference (pp)",
        "Risk ratio",
        "OR (95% CI)",
        "P value",
    ]
    body = [
        [
            r["contrast"].replace("h", " h"),
            r["reference_30d_risk"].replace("%", ""),
            r["comparison_30d_risk"].replace("%", ""),
            r["risk_difference"].replace(" pp", ""),
            r["risk_ratio"],
            r["odds_ratio_95_ci"],
            r["p_value"],
        ]
        for r in rows
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False)

    # Sum is 9360 DXA, the usable width of a Letter page with 1-inch margins.
    widths = [1780, 1260, 1390, 1390, 940, 1660, 940]
    set_table_layout(table, widths)

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(table.rows[0].cells[i], bottom=True)

    for data_row in body:
        cells = table.add_row().cells
        set_table_layout(table, widths)
        for i, value in enumerate(data_row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, bold=False, align=align)

    # Reapply three-line table geometry after adding rows.
    table_borders(table, top=True, bottom=True, inside_h=False, inside_v=False)
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    note_run = note.add_run(
        "Note: Risks are model-standardized 30-day mortality estimates from the "
        "IPCW-weighted, covariate-adjusted pooled logistic regression model. "
        "The reference and comparison risks correspond to the two strategies named "
        "in each contrast. Risk differences are shown in percentage points. "
        "CI, confidence interval; IPCW, inverse probability of censoring weighting; "
        "MIMIC-IV, Medical Information Mart for Intensive Care IV; OR, odds ratio; "
        "pp, percentage points."
    )
    note_run.font.name = "Times New Roman"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    note_run.font.size = Pt(9)

    # Clear accidental spacing in all table paragraphs.
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

    outputs = [WORK_OUT] + ([EXTRA_OUT] if EXTRA_OUT else [])
    for out in outputs:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)

    for out in outputs:
        print(out)


if __name__ == "__main__":
    main()

